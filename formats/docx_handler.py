import os
import logging
from typing import Dict, Any
from .base_handler import BaseFormatHandler

logger = logging.getLogger(__name__)

class DOCXHandler(BaseFormatHandler):
    """Обработчик для DOCX файлов"""
    
    @staticmethod
    def can_handle(file_path: str) -> bool:
        ext = BaseFormatHandler.get_file_extension(file_path)
        return ext in ['.docx', '.doc']
    
    @staticmethod
    def extract_text(file_path: str, parameters: Dict[str, Any]) -> str:
        action_type = parameters.get('type', 'first_chars')
        amount = parameters.get('amount', 500)
        
        try:
            import docx
            from docx import Document
            
            doc = Document(file_path)
            full_text = []
            
            # Извлекаем текст из параграфов
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text)
            
            # Извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text.strip():
                                full_text.append(paragraph.text)
            
            combined_text = "\n".join(full_text)
            
            if action_type == 'first_chars':
                return combined_text[:amount]
            else:
                return combined_text
                
        except ImportError:
            return "Ошибка: python-docx не установлен. Установите: pip install python-docx"
        except Exception as e:
            logger.error(f"Ошибка при обработке DOCX {file_path}: {e}")
            return f"Ошибка при обработке DOCX: {str(e)}"
    
    @staticmethod
    def get_metadata(file_path: str) -> Dict[str, str]:
        """Извлекает метаданные из DOCX"""
        try:
            import docx
            from docx import Document
            
            doc = Document(file_path)
            core_properties = doc.core_properties
            
            metadata = {
                'title': core_properties.title,
                'author': core_properties.author,
                'subject': core_properties.subject,
                'keywords': core_properties.keywords,
                'comments': core_properties.comments,
                'last_modified_by': core_properties.last_modified_by,
                'created': str(core_properties.created) if core_properties.created else '',
                'modified': str(core_properties.modified) if core_properties.modified else '',
                'category': core_properties.category,
                'version': core_properties.version
            }
            
            return {k: v for k, v in metadata.items() if v}
            
        except Exception as e:
            logger.error(f"Ошибка при извлечении метаданных DOCX: {e}")
            return {}
